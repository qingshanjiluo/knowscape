"""
KnowScape Cloud Document Parser Engine
=======================================
Python implementation for cloud (FastAPI) deployment.
Mirrors the Rust knowscape-parser API.

Usage:
    from parsers import ParserEngine
    engine = ParserEngine()
    doc = engine.parse("book.epub")
    print(doc.metadata.title)
    print(doc.full_text[:500])
"""

import hashlib
import json
import logging
import os
import re
import subprocess
import tempfile
import uuid
from dataclasses import dataclass, field, asdict
from datetime import datetime, timezone
from enum import Enum
from pathlib import Path
from typing import Optional, List
from urllib.parse import urlparse

logger = logging.getLogger("knowscape.parser")


# ─── Types ────────────────────────────────────────────────────

class DocumentFormat(str, Enum):
    EPUB = "epub"
    PDF = "pdf"
    MOBI = "mobi"
    AZW = "azw"
    LATEX = "latex"
    DOCX = "docx"
    MARKDOWN = "markdown"
    HTML = "html"
    URL = "url"
    PLAIN_TEXT = "plain_text"
    UNKNOWN = "unknown"

    @classmethod
    def from_path(cls, path: str) -> "DocumentFormat":
        lower = path.lower()
        if lower.startswith("http://") or lower.startswith("https://"):
            return cls.URL
        ext = Path(lower).suffix.lstrip(".")
        mapping = {
            "epub": cls.EPUB,
            "pdf": cls.PDF,
            "mobi": cls.MOBI,
            "azw": cls.AZW, "azw3": cls.AZW, "azw4": cls.AZW,
            "tex": cls.LATEX, "ltx": cls.LATEX,
            "docx": cls.DOCX,
            "md": cls.MARKDOWN, "markdown": cls.MARKDOWN,
            "htm": cls.HTML, "html": cls.HTML, "xhtml": cls.HTML,
            "txt": cls.PLAIN_TEXT,
        }
        return mapping.get(ext, cls.UNKNOWN)

    def description(self) -> str:
        return {
            self.EPUB: "EPUB eBook",
            self.PDF: "PDF Document",
            self.MOBI: "MOBI eBook",
            self.AZW: "AZW eBook",
            self.LATEX: "LaTeX Document",
            self.DOCX: "Word Document",
            self.MARKDOWN: "Markdown",
            self.HTML: "HTML",
            self.URL: "Web URL",
            self.PLAIN_TEXT: "Plain Text",
            self.UNKNOWN: "Unknown Format",
        }.get(self, "Unknown")


@dataclass
class Chapter:
    index: int
    title: str
    level: int
    content: str
    start_pos: int
    end_pos: int


@dataclass
class DocumentMetadata:
    title: str
    author: str
    language: str
    source_format: DocumentFormat
    source_path: str
    source_hash: str
    total_chapters: int
    word_count: int
    ocr_applied: bool


@dataclass
class ParsedDocument:
    id: str
    metadata: DocumentMetadata
    full_text: str
    chapters: List[Chapter]
    notes: List[str]

    def to_json(self) -> str:
        return json.dumps(self._to_dict(), ensure_ascii=False, indent=2)

    def _to_dict(self):
        return {
            "id": self.id,
            "metadata": asdict(self.metadata),
            "full_text": self.full_text,
            "chapters": [asdict(c) for c in self.chapters],
            "notes": self.notes,
        }

    @classmethod
    def from_json(cls, data: str) -> "ParsedDocument":
        d = json.loads(data)
        md = DocumentMetadata(**d["metadata"])
        md.source_format = DocumentFormat(md.source_format)
        chapters = [Chapter(**c) for c in d["chapters"]]
        return cls(
            id=d["id"],
            metadata=md,
            full_text=d["full_text"],
            chapters=chapters,
            notes=d.get("notes", []),
        )


@dataclass
class ParseOptions:
    extract_chapters: bool = True
    preserve_formatting: bool = True
    ocr_language: str = "chi_sim+eng"
    ocr_enabled: bool = True
    ocr_dpi: int = 300


# ─── Utilities ────────────────────────────────────────────────

def hash_file(path: str) -> str:
    h = hashlib.sha256()
    with open(path, "rb") as f:
        for chunk in iter(lambda: f.read(65536), b""):
            h.update(chunk)
    return h.hexdigest()


def hash_text(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8")).hexdigest()


def normalize_text(text: str) -> str:
    text = text.strip()
    text = text.removeprefix("\ufeff")
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    return text


def is_chinese_text(text: str) -> bool:
    chinese_count = sum(1 for c in text if '\u4e00' <= c <= '\u9fff')
    return chinese_count > len(text) * 0.1 if text else False


def extract_title_from_markdown(text: str) -> Optional[str]:
    for line in text.splitlines():
        stripped = line.strip()
        if stripped.startswith("# "):
            return stripped[2:].strip()
    return None


def split_into_chapters(full_text: str) -> List[tuple]:
    """Split text into chapters by H1/H2 headings."""
    chapters = []
    current_title = "前言"
    current_level = 1
    current_start = 0
    index = 0

    heading_re = re.compile(r"^(#{1,6})\s+(.+)$", re.MULTILINE)

    for match in heading_re.finditer(full_text):
        level = len(match.group(1))
        title = match.group(2).strip()
        pos = match.start()

        if pos > current_start or index > 0:
            chapters.append((index, current_title, current_level, current_start, pos))
            index += 1

        current_title = title
        current_level = level
        current_start = pos

    # Last chapter
    end = len(full_text)
    if not chapters:
        chapters.append((0, current_title, current_level, 0, end))
    else:
        chapters.append((index, current_title, current_level, current_start, end))

    return chapters


def detect_and_decode(data: bytes) -> str:
    """Detect encoding and decode bytes to string."""
    try:
        return data.decode("utf-8")
    except UnicodeDecodeError:
        try:
            return data.decode("gbk")
        except UnicodeDecodeError:
            return data.decode("utf-8", errors="replace")


# ─── Individual Parsers ──────────────────────────────────────

class EpubParser:
    """Parse EPUB files using Python's ebooklib."""

    @staticmethod
    def parse(path: str, options: ParseOptions) -> ParsedDocument:
        import ebooklib
        from ebooklib import epub
        from bs4 import BeautifulSoup

        logger.info("Parsing EPUB: %s", path)
        source_hash = hash_file(path)

        book = epub.read_epub(path)

        # Metadata
        title = book.get_metadata("DC", "title")
        title = title[0][0] if title else Path(path).stem

        author = book.get_metadata("DC", "creator")
        author = author[0][0] if author else ""

        language = book.get_metadata("DC", "language")
        language = language[0][0] if language else "zh"

        # Extract chapters from spine
        chapters = []
        full_texts = []
        notes = []

        for i, item in enumerate(book.get_items_of_type(ebooklib.ITEM_DOCUMENT)):
            content = item.get_body_content().decode("utf-8", errors="replace")
            soup = BeautifulSoup(content, "html.parser")

            # Remove scripts and styles
            for tag in soup(["script", "style", "noscript"]):
                tag.decompose()

            # Extract title from first heading
            h1 = soup.find("h1")
            chap_title = h1.get_text(strip=True) if h1 else f"第{i+1}章"

            # Convert to simple markdown
            md = EpubParser._soup_to_markdown(soup)

            if not md.strip():
                continue

            start = sum(len(t) for t in full_texts) + len(full_texts)
            full_texts.append(md)
            end = start + len(md)

            chapters.append(Chapter(
                index=i, title=chap_title, level=1,
                content=md, start_pos=start, end_pos=end,
            ))

        if not chapters:
            raise ValueError("Empty document")

        full_text = "\n\n".join(full_texts)

        return ParsedDocument(
            id=str(uuid.uuid4()),
            metadata=DocumentMetadata(
                title=title, author=author, language=language,
                source_format=DocumentFormat.EPUB,
                source_path=path, source_hash=source_hash,
                total_chapters=len(chapters),
                word_count=len(full_text), ocr_applied=False,
            ),
            full_text=full_text, chapters=chapters, notes=notes,
        )

    @staticmethod
    def _soup_to_markdown(soup) -> str:
        lines = []
        for tag in soup.find_all(["h1", "h2", "h3", "h4", "h5", "h6", "p", "li", "blockquote", "pre"]):
            text = tag.get_text(strip=True)
            if not text:
                continue
            tagname = tag.name
            if tagname.startswith("h"):
                level = tagname[1]
                lines.append(f"{'#' * int(level)} {text}\n")
            elif tagname == "p":
                lines.append(f"{text}\n")
            elif tagname == "li":
                lines.append(f"- {text}")
            elif tagname == "blockquote":
                lines.append(f"> {text}\n")
            elif tagname == "pre":
                lines.append(f"```\n{text}\n```\n")
        return "\n".join(lines)


class PdfParser:
    """Parse PDF files. Uses PyMuPDF for text PDFs, Tesseract for scanned."""

    @staticmethod
    def parse(path: str, options: ParseOptions) -> ParsedDocument:
        logger.info("Parsing PDF: %s", path)
        source_hash = hash_file(path)

        import fitz  # PyMuPDF

        doc = fitz.open(path)
        total_pages = len(doc)

        # Try text extraction first
        all_text = []
        text_page_count = 0

        for page_num in range(total_pages):
            page = doc[page_num]
            text = page.get_text().strip()
            if len(text) > 50:
                text_page_count += 1
            all_text.append(text)

        doc.close()

        full_text = "\n\n".join(all_text)
        has_meaningful_text = text_page_count > total_pages * 0.3

        if has_meaningful_text:
            logger.info("PDF text extraction successful (%d/%d pages)", text_page_count, total_pages)
            return PdfParser._build_result(path, source_hash, full_text, ocr_applied=False)
        else:
            logger.warning("PDF has little text (%d/%d pages), trying OCR", text_page_count, total_pages)
            if options.ocr_enabled:
                return PdfParser._ocr_pdf(path, source_hash, options)
            else:
                return PdfParser._build_result(path, source_hash, full_text, ocr_applied=False)

    @staticmethod
    def _build_result(path: str, source_hash: str, full_text: str, ocr_applied: bool) -> ParsedDocument:
        normalized = normalize_text(full_text)
        title = extract_title_from_markdown(normalized) or Path(path).stem
        chapter_data = split_into_chapters(normalized)
        chapters = []

        if not chapter_data:
            chapters.append(Chapter(0, title, 1, normalized, 0, len(normalized)))
        else:
            for idx, ct, lv, st, en in chapter_data:
                chapters.append(Chapter(idx, ct, lv, normalized[st:en].strip(), st, en))

        return ParsedDocument(
            id=str(uuid.uuid4()),
            metadata=DocumentMetadata(
                title=title, author="", language="zh",
                source_format=DocumentFormat.PDF,
                source_path=path, source_hash=source_hash,
                total_chapters=len(chapters),
                word_count=len(normalized), ocr_applied=ocr_applied,
            ),
            full_text=normalized, chapters=chapters,
            notes=[f"OCR applied" if ocr_applied else "Text-based PDF"],
        )

    @staticmethod
    def _ocr_pdf(path: str, source_hash: str, options: ParseOptions) -> ParsedDocument:
        import fitz
        import pytesseract
        from PIL import Image
        import io

        logger.info("OCR processing PDF: %s", path)
        doc = fitz.open(path)
        all_text = []
        notes = []

        for page_num in range(len(doc)):
            page = doc[page_num]
            pix = page.get_pixmap(dpi=options.ocr_dpi)
            img_data = pix.tobytes("png")
            img = Image.open(io.BytesIO(img_data))
            try:
                text = pytesseract.image_to_string(img, lang=options.ocr_language)
                all_text.append(text)
            except Exception as e:
                notes.append(f"OCR failed for page {page_num+1}: {e}")
                all_text.append("")

        doc.close()
        full_text = "\n\n".join(all_text)
        notes.append(f"OCR applied on {len(doc)} pages")
        result = PdfParser._build_result(path, source_hash, full_text, ocr_applied=True)
        result.notes = notes
        return result


class MobiParser:
    """Parse MOBI/AZW via Calibre ebook-convert."""

    @staticmethod
    def parse(path: str, options: ParseOptions) -> ParsedDocument:
        logger.info("Parsing MOBI: %s", path)
        source_hash = hash_file(path)

        with tempfile.NamedTemporaryFile(suffix=".md", delete=False) as tmp:
            output_path = tmp.name

        try:
            result = subprocess.run(
                ["ebook-convert", path, output_path, "--input-encoding", "utf-8"],
                capture_output=True, text=True, timeout=120,
            )
            if result.returncode != 0:
                raise RuntimeError(f"ebook-convert failed: {result.stderr}")

            with open(output_path, "r", encoding="utf-8", errors="replace") as f:
                markdown = f.read()
        except FileNotFoundError:
            raise RuntimeError(
                "Calibre (ebook-convert) not found. Install from https://calibre-ebook.com"
            )
        finally:
            if os.path.exists(output_path):
                os.unlink(output_path)

        normalized = normalize_text(markdown)
        title = extract_title_from_markdown(normalized) or Path(path).stem

        chapter_data = split_into_chapters(normalized)
        chapters = []
        for idx, ct, lv, st, en in chapter_data:
            chapters.append(Chapter(idx, ct, lv, normalized[st:en].strip(), st, en))

        return ParsedDocument(
            id=str(uuid.uuid4()),
            metadata=DocumentMetadata(
                title=title, author="", language="zh",
                source_format=DocumentFormat(path.split(".")[-1]),
                source_path=path, source_hash=source_hash,
                total_chapters=len(chapters),
                word_count=len(normalized), ocr_applied=False,
            ),
            full_text=normalized, chapters=chapters,
            notes=["Converted via Calibre (ebook-convert)"],
        )


class LatexParser:
    """Parse LaTeX via Pandoc."""

    @staticmethod
    def parse(path: str, options: ParseOptions) -> ParsedDocument:
        logger.info("Parsing LaTeX: %s", path)
        source_hash = hash_file(path)

        with tempfile.NamedTemporaryFile(suffix=".md", delete=False) as tmp:
            output_path = tmp.name

        try:
            result = subprocess.run(
                ["pandoc", path, "-f", "latex", "-t", "markdown",
                 "--wrap=preserve", "-o", output_path],
                capture_output=True, text=True, timeout=60,
            )
            if result.returncode != 0:
                raise RuntimeError(f"Pandoc failed: {result.stderr}")

            with open(output_path, "r", encoding="utf-8", errors="replace") as f:
                markdown = f.read()
        except FileNotFoundError:
            raise RuntimeError("Pandoc not found. Install from https://pandoc.org")
        finally:
            if os.path.exists(output_path):
                os.unlink(output_path)

        normalized = normalize_text(markdown)
        title = extract_title_from_markdown(normalized) or Path(path).stem
        chapter_data = split_into_chapters(normalized)
        chapters = []
        for idx, ct, lv, st, en in chapter_data:
            chapters.append(Chapter(idx, ct, lv, normalized[st:en].strip(), st, en))

        return ParsedDocument(
            id=str(uuid.uuid4()),
            metadata=DocumentMetadata(
                title=title, author="", language="zh",
                source_format=DocumentFormat.LATEX,
                source_path=path, source_hash=source_hash,
                total_chapters=len(chapters), word_count=len(normalized),
                ocr_applied=False,
            ),
            full_text=normalized, chapters=chapters,
            notes=["Converted via Pandoc"],
        )


class DocxParser:
    """Parse DOCX files using python-docx."""

    @staticmethod
    def parse(path: str, options: ParseOptions) -> ParsedDocument:
        from docx import Document

        logger.info("Parsing DOCX: %s", path)
        source_hash = hash_file(path)

        doc = Document(path)
        full_text = []
        chapters = []
        current_title = "正文"
        current_level = 1
        chapter_buf = []
        chapter_index = 0
        char_count = 0

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            style = para.style.name if para.style else ""

            if style.startswith("Heading"):
                # Save previous chapter
                if chapter_buf:
                    content = "\n\n".join(chapter_buf)
                    chapters.append(Chapter(
                        chapter_index, current_title, current_level,
                        content, char_count - len(content), char_count,
                    ))
                    chapter_index += 1
                    chapter_buf = []

                level = int(style[-1]) if style[-1].isdigit() else 1
                current_title = text
                current_level = level
                prefix = "#" * level
                full_text.append(f"{prefix} {text}\n")
                chapter_buf.append(f"{prefix} {text}")
            else:
                full_text.append(f"{text}\n")
                chapter_buf.append(text)

            char_count += len(text) + 1

        # Last chapter
        if chapter_buf:
            content = "\n\n".join(chapter_buf)
            chapters.append(Chapter(
                chapter_index, current_title, current_level,
                content, char_count - len(content), char_count,
            ))

        if not chapters:
            chapters.append(Chapter(0, "正文", 1, "\n".join(full_text), 0, char_count))

        full = "\n".join(full_text)
        title = extract_title_from_markdown(full) or Path(path).stem

        return ParsedDocument(
            id=str(uuid.uuid4()),
            metadata=DocumentMetadata(
                title=title,
                author=doc.core_properties.author or "",
                language="zh",
                source_format=DocumentFormat.DOCX,
                source_path=path, source_hash=source_hash,
                total_chapters=len(chapters), word_count=len(full),
                ocr_applied=False,
            ),
            full_text=full, chapters=chapters, notes=[],
        )


class HtmlParser:
    """Parse HTML files or URLs. Extracts readable content."""

    @staticmethod
    def parse(path_or_url: str, options: ParseOptions) -> ParsedDocument:
        import requests
        from bs4 import BeautifulSoup

        is_url = path_or_url.startswith("http://") or path_or_url.startswith("https://")

        if is_url:
            logger.info("Fetching URL: %s", path_or_url)
            resp = requests.get(
                path_or_url,
                headers={"User-Agent": "KnowScape/0.1"},
                timeout=30,
            )
            resp.raise_for_status()
            html_content = resp.text
            source_hash = hash_text(html_content)
        else:
            logger.info("Parsing HTML file: %s", path_or_url)
            with open(path_or_url, "rb") as f:
                raw = f.read()
            html_content = detect_and_decode(raw)
            source_hash = hash_file(path_or_url)

        soup = BeautifulSoup(html_content, "html.parser")

        # Remove unwanted elements
        for tag in soup(["script", "style", "noscript", "iframe", "nav", "footer", "header", "aside"]):
            tag.decompose()

        # Extract title
        title = ""
        if soup.title:
            title = soup.title.get_text(strip=True)
        if not title:
            og_title = soup.find("meta", property="og:title")
            if og_title and og_title.get("content"):
                title = og_title["content"]
        if not title:
            h1 = soup.find("h1")
            if h1:
                title = h1.get_text(strip=True)
        if not title:
            title = Path(path_or_url).stem if not is_url else urlparse(path_or_url).netloc

        # Extract author
        author = ""
        meta_author = soup.find("meta", attrs={"name": "author"})
        if meta_author and meta_author.get("content"):
            author = meta_author["content"]

        # Find main content
        main_content = None
        for sel in ["main", "article", '[role="main"]', ".post-content",
                     ".article-content", ".entry-content", ".content", "#content"]:
            main_content = soup.select_one(sel)
            if main_content:
                break

        if not main_content:
            main_content = soup.find("body") or soup

        # Convert to Markdown
        md = HtmlParser._soup_to_markdown(main_content)
        normalized = normalize_text(md)

        if not normalized:
            normalized = soup.get_text(strip=True)

        chapter_data = split_into_chapters(normalized)
        chapters = []
        for idx, ct, lv, st, en in chapter_data:
            chapters.append(Chapter(idx, ct, lv, normalized[st:en].strip(), st, en))

        notes = [f"Fetched from {path_or_url}"] if is_url else []

        return ParsedDocument(
            id=str(uuid.uuid4()),
            metadata=DocumentMetadata(
                title=title, author=author,
                language="zh",
                source_format=DocumentFormat.URL if is_url else DocumentFormat.HTML,
                source_path=path_or_url, source_hash=source_hash,
                total_chapters=len(chapters), word_count=len(normalized),
                ocr_applied=False,
            ),
            full_text=normalized, chapters=chapters, notes=notes,
        )

    @staticmethod
    def _soup_to_markdown(soup) -> str:
        lines = []
        for tag in soup.find_all(["h1", "h2", "h3", "h4", "h5", "h6", "p", "li", "blockquote", "pre", "hr"]):
            text = tag.get_text(strip=True)
            tn = tag.name
            if tn.startswith("h"):
                lines.append(f"{'#' * int(tn[1])} {text}\n")
            elif tn == "p":
                # Process links within paragraph
                inline = HtmlParser._inline_to_markdown(tag)
                if inline.strip():
                    lines.append(f"{inline}\n")
            elif tn == "li":
                lines.append(f"- {text}")
            elif tn == "blockquote":
                for qline in text.splitlines():
                    lines.append(f"> {qline}")
                lines.append("")
            elif tn == "pre":
                lines.append(f"```\n{text}\n```\n")
            elif tn == "hr":
                lines.append("---\n")
        return "\n".join(lines)

    @staticmethod
    def _inline_to_markdown(tag) -> str:
        result = ""
        for child in tag.children:
            if hasattr(child, "name"):
                inner = child.get_text(strip=True)
                if child.name == "a":
                    href = child.get("href", "")
                    result += f"[{inner}]({href})"
                elif child.name in ("strong", "b"):
                    result += f"**{inner}**"
                elif child.name in ("em", "i"):
                    result += f"*{inner}*"
                elif child.name == "code":
                    result += f"`{inner}`"
                elif child.name == "br":
                    result += "\n"
                else:
                    result += inner
            else:
                result += str(child).strip()
        return result


class MarkdownParser:
    """Parse Markdown files — validate and split into chapters."""

    @staticmethod
    def parse(path: str, options: ParseOptions) -> ParsedDocument:
        logger.info("Parsing Markdown: %s", path)
        with open(path, "rb") as f:
            raw = f.read()
        content = detect_and_decode(raw)
        return MarkdownParser._from_string(content, path, hash_file(path))

    @staticmethod
    def parse_from_string(content: str, filename: str = "untitled.md") -> ParsedDocument:
        logger.info("Parsing Markdown from string: %d chars", len(content))
        return MarkdownParser._from_string(content, filename, hash_text(content))

    @staticmethod
    def _from_string(content: str, path: str, source_hash: str) -> ParsedDocument:
        normalized = normalize_text(content)
        if not normalized.strip():
            raise ValueError("Empty document")

        title = extract_title_from_markdown(normalized) or Path(path).stem

        chapter_data = split_into_chapters(normalized)
        chapters = []
        if not chapter_data:
            chapters.append(Chapter(0, title, 1, normalized, 0, len(normalized)))
        else:
            for idx, ct, lv, st, en in chapter_data:
                chapters.append(Chapter(idx, ct, lv, normalized[st:en].strip(), st, en))

        return ParsedDocument(
            id=str(uuid.uuid4()),
            metadata=DocumentMetadata(
                title=title, author="",
                language="zh" if is_chinese_text(normalized) else "en",
                source_format=DocumentFormat.MARKDOWN,
                source_path=path, source_hash=source_hash,
                total_chapters=len(chapters), word_count=len(normalized),
                ocr_applied=False,
            ),
            full_text=normalized, chapters=chapters,
            notes=[] if path != "untitled.md" else ["Manually uploaded Markdown"],
        )


# ─── Parser Engine (Facade) ──────────────────────────────────

class ParserEngine:
    """Main entry point — detects format and dispatches to the appropriate parser."""

    def __init__(self):
        self._parsers = {
            DocumentFormat.EPUB: EpubParser,
            DocumentFormat.PDF: PdfParser,
            DocumentFormat.MOBI: MobiParser,
            DocumentFormat.AZW: MobiParser,
            DocumentFormat.LATEX: LatexParser,
            DocumentFormat.DOCX: DocxParser,
            DocumentFormat.MARKDOWN: MarkdownParser,
            DocumentFormat.PLAIN_TEXT: MarkdownParser,
            DocumentFormat.HTML: HtmlParser,
            DocumentFormat.URL: HtmlParser,
        }

    def parse(self, path_or_url: str, options: Optional[ParseOptions] = None) -> ParsedDocument:
        if options is None:
            options = ParseOptions()
        fmt = DocumentFormat.from_path(path_or_url)

        logger.info("Detected format: %s for %s", fmt.value, path_or_url)

        parser_cls = self._parsers.get(fmt)
        if parser_cls is None:
            # Fallback: try as plain text
            if os.path.exists(path_or_url):
                with open(path_or_url, "rb") as f:
                    raw = f.read()
                text = detect_and_decode(raw)
                return MarkdownParser.parse_from_string(text, path_or_url)
            raise ValueError(f"Unsupported format: {path_or_url}")

        return parser_cls.parse(path_or_url, options)

    def parse_with_format(
        self, path: str, fmt: DocumentFormat,
        options: Optional[ParseOptions] = None,
    ) -> ParsedDocument:
        if options is None:
            options = ParseOptions()
        parser_cls = self._parsers.get(fmt)
        if parser_cls is None:
            raise ValueError(f"Unsupported format: {fmt}")
        return parser_cls.parse(path, options)

    @staticmethod
    def supported_formats() -> List[tuple]:
        return [
            (DocumentFormat.EPUB, "EPUB eBook"),
            (DocumentFormat.PDF, "PDF Document"),
            (DocumentFormat.MOBI, "MOBI eBook (requires Calibre)"),
            (DocumentFormat.AZW, "AZW eBook (requires Calibre)"),
            (DocumentFormat.LATEX, "LaTeX Document (requires Pandoc)"),
            (DocumentFormat.DOCX, "Word Document"),
            (DocumentFormat.MARKDOWN, "Markdown"),
            (DocumentFormat.HTML, "HTML File"),
            (DocumentFormat.URL, "Web URL"),
        ]


# ─── Requirements ─────────────────────────────────────────────
# Install with:
#   pip install ebooklib beautifulsoup4 PyMuPDF pytesseract Pillow
#   pip install python-docx requests lxml
#   # Optional: pip install pandoc (or install Pandoc system-wide)
#   # Optional: pip install calibre (or install Calibre system-wide)
